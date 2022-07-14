import docx
import sys
from PIL import Image
from PyQt5 import QtWidgets, QtCore
from PyQt5.QtWidgets import QDialog, QApplication, QFileDialog, QLabel, QMessageBox
from PyQt5.uic import loadUi
import source
from timeit import default_timer as timer

class Ui_Dialog(object):
    def setupUi(self, Dialog):
        Dialog.setObjectName("Dialog")
        Dialog.resize(800, 600)
        Dialog.setStyleSheet("background-color: rgb(236, 236, 236);")
        self.encrypt = QtWidgets.QPushButton(Dialog)
        self.encrypt.setGeometry(QtCore.QRect(230, 450, 171, 51))
        self.encrypt.setStyleSheet("background-color: qlineargradient(spread:pad, x1:0.0395455, y1:0.034, x2:1, y2:1, stop:0.193182 rgba(140, 160, 210, 255), stop:1 rgba(137, 255, 205, 255));\n"
"font: 57 14pt \"Poppins Medium\"\n"
"rgb(85, 85, 255)\n"
"")
        self.encrypt.setObjectName("encrypt")
        self.decrypt = QtWidgets.QPushButton(Dialog)
        self.decrypt.setGeometry(QtCore.QRect(440, 450, 171, 51))
        self.decrypt.setStyleSheet("background-color: qlineargradient(spread:pad, x1:0.0395455, y1:0.034, x2:1, y2:1, stop:0.193182 rgba(140, 160, 210, 255), stop:1 rgba(137, 255, 205, 255));\n"
"font: 57 14pt \"Poppins Medium\"\n"
"rgb(85, 85, 255)")
        self.decrypt.setObjectName("decrypt")
        self.Title = QtWidgets.QLabel(Dialog)
        self.Title.setGeometry(QtCore.QRect(0, 0, 801, 121))
        self.Title.setStyleSheet("background-color: qlineargradient(spread:pad, x1:0.0395455, y1:0.034, x2:1, y2:1, stop:0.193182 rgba(140, 160, 210, 255), stop:1 rgba(137, 255, 205, 255));\n"
"font: 57 16pt \"Poppins Medium\";\n"
"color: rgb(255, 255, 255);")
        self.Title.setAlignment(QtCore.Qt.AlignCenter)
        self.Title.setObjectName("Title")
        self.label = QtWidgets.QLabel(Dialog)
        self.label.setGeometry(QtCore.QRect(190, 170, 451, 271))
        self.label.setStyleSheet("image: url(:/newPrefix/image.png);")
        self.label.setText("")
        self.label.setObjectName("label")

        self.retranslateUi(Dialog)
        QtCore.QMetaObject.connectSlotsByName(Dialog)

    def retranslateUi(self, Dialog):
        _translate = QtCore.QCoreApplication.translate
        Dialog.setWindowTitle(_translate("Dialog", "Dialog"))
        self.encrypt.setText(_translate("Dialog", "Enkripsi"))
        self.decrypt.setText(_translate("Dialog", "Dekripsi"))
        self.Title.setText(_translate("Dialog", "Pengamanan Data Dengan Gronsfeld Cipher\n"
"dan Steganografi End Of File"))

class MainWindow(QDialog):
    def __init__(self):
        super(MainWindow,self).__init__()
        loadUi("main.ui",self)
        self.decrypt.clicked.connect(self.extractpage)
        self.encrypt.clicked.connect(self.encryptpage)
        self.keluar.clicked.connect(self.exit)
    
    def exit(self):
        msg = QMessageBox.question(self, "Exit", "Apakah kamu yakin ingin keluar?", QMessageBox.No|QMessageBox.Yes)
        if (msg == QMessageBox.Yes):
            QApplication.closeAllWindows()

    def extractpage(self):
        screen2=Decrypt()
        widget.addWidget(screen2)
        widget.setCurrentIndex(widget.currentIndex()+1)

    def encryptpage(self):
        screen3=Encrypt()
        widget.addWidget(screen3)
        widget.setCurrentIndex(widget.currentIndex()+1)

class Decrypt(QDialog):
    def __init__(self):
        super(Decrypt,self).__init__()
        loadUi("decrypt.ui",self)
        self.data_global = ''
        self.text_global = ''
        self.decrypt_global = ''
        self.browse2.clicked.connect(self.browseimage)
        self.save.clicked.connect(self.cari)
        self.proses.clicked.connect(self.vincky)
        self.Back.clicked.connect(self.mainpage)
    
    def mainpage(self):
        screen1=MainWindow()
        widget.addWidget(screen1)
        widget.setCurrentIndex(widget.currentIndex()+1)
    
    def browseimage(self):
        fname=QFileDialog.getOpenFileName(self, "Select Image", 'Done/', "Image (*.png)")
        #open the image
        self.image.setText(fname[0])
        self.image_global = str(fname[0])

    def decode(self):
        image = Image.open(self.image_global, 'r')

        self.data_global = ''
        imgdata = iter(image.getdata())

        while (True):
            pixels = [value for value in imgdata.__next__()[:3] + imgdata.__next__()[:3] + imgdata.__next__()[:3]]

            # string of binary data
            binstr = ''

            for i in pixels[:8]:
                if (i % 2 == 0):
                    binstr += '0'
                else:
                    binstr += '1'

            self.data_global += chr(int(binstr, 2))
            if (pixels[-1] % 2 != 0):
                return self.data_global
    
    def decrypt(self):
        kata = self.data_global
        value = [ord(x) for x in kata]
        key = [int(x) for x in self.key.text()]
        
        temp = 0
        new_key = []
        for i in range (len(value)):
            new_key.append(key[temp])
            if temp <= len(key)-2:
                temp = temp + 1
            else:
                temp = 0

        decrypt = [(value[i]-new_key[i]) for i in range (len(value))]

        for e in decrypt:
            self.decrypt_global = self.decrypt_global + chr(e)

    def cari(self):
        self.output = QFileDialog.getSaveFileName(None, "Save file", "Done/", "Text Files (*.txt)")[0]
        self.outputname.setText(self.output)
        self.text_global = self.output
    
    def vincky(self):
        self.decode()
        self.decrypt()
        start = timer()

        x = self.text_global.split('/')[-1].split('.')[-1]
        if (x == 'docx'):
            data = docx.Document()
            data.add_paragraph(self.decrypt_global)
            data.save(self.output)
        elif (x == 'txt'):
            data = self.decrypt_global
            with open(self.output, 'w') as f:
                f.write(data)
        
        end = timer()
        elapsed_time = end - start
        print(f"[INFO] Waktu diperlukan pada proses ekstraksi adalah {elapsed_time} s")
        msg = QMessageBox.question(self, "Ekstraksi Stego-Iameg Berhasil", "Apakah kamu ingin mengekstraksi stego-image lagi?", QMessageBox.No|QMessageBox.Yes)
        if (msg == QMessageBox.No):
            QApplication.closeAllWindows()
        elif(msg == QMessageBox.Yes):
            screen=Decrypt()
            widget.addWidget(screen)
            widget.setCurrentIndex(widget.currentIndex()+1)

class Encrypt(QDialog):
    def __init__(self):
        super(Encrypt,self).__init__()
        loadUi("encrypt.ui",self)
        self.fname_global = ''
        self.image_global = ''
        self.encrypt_global = ''
        self.data_global = ''
        self.browse.clicked.connect(self.browsefile)
        self.browse2.clicked.connect(self.browseimage)
        self.proses.clicked.connect(self.vincky)
        self.save.clicked.connect(self.cari)
        self.Back.clicked.connect(self.mainpage)

    def mainpage(self):
        screen1=MainWindow()
        widget.addWidget(screen1)
        widget.setCurrentIndex(widget.currentIndex()+1)

    def browsefile(self):
        fname = QFileDialog.getOpenFileName(self, "Open File", "Document/", "Text File (*.txt)")
        self.filename.setText(fname[0])
        self.fname_global = str(fname[0])

    def ReadingText(self):
        x = self.fname_global.split('/')[-1].split('.')[-1]
        if (x == 'docx'):
            doc = docx.Document(self.fname_global)
            completedText = []
            for paragraph in doc.paragraphs:
                completedText.append(paragraph.text)
            return '\n' .join(completedText)
        elif (x == 'txt'):
            with open(self.fname_global) as f:
                lines = f.readlines()
            return '' .join(lines)
    
    def Encrypt(self):
        kata = self.ReadingText()
        value = [ord(x) for x in kata]
        key = [int(x) for x in self.key.text()]
        
        temp = 0
        new_key = []
        for i in range (len(value)):
            new_key.append(key[temp])
            if temp <= len(key)-2:
                temp = temp + 1
            else:
                temp = 0

        encrypt = [(value[i]+new_key[i]) for i in range (len(value))]

        for e in encrypt:
            self.encrypt_global = self.encrypt_global + chr(e)

    def browseimage(self):
        fname=QFileDialog.getOpenFileName(self, "Select Image", "Gambar/", "Image (*.png)")
        #open the image
        self.image.setText(fname[0])
        self.image_global = str(fname[0])
    
    def genData(self,data):
		# list of binary codes of given data 
        newd = []
        for i in data:
            newd.append(format(ord(i), '08b'))
        return newd

    # Pixels are modified according to the 8-bit binary data and finally returned
    def modPix(self, pix, data):
        datalist = self.genData(data)
        lendata = len(datalist)
        imdata = iter(pix)
        for i in range(lendata):
            # Extracting 3 pixels at a time
            pix = [value for value in imdata.__next__()[:3] + imdata.__next__()[:3] + imdata.__next__()[:3]]

            # Pixel value should be made odd for 1 and even for 0
            for j in range(0, 8):
                if (datalist[i][j] == '0' and pix[j]% 2 != 0):
                    pix[j] -= 1

                elif (datalist[i][j] == '1' and pix[j] % 2 == 0):
                    if(pix[j] != 0):
                        pix[j] -= 1
                    else:
                        pix[j] += 1

            # Eighth pixel of every set tells whether to stop ot read further. 0 means keep reading; 1 means thec message is over.
            if (i == lendata - 1):
                if (pix[-1] % 2 == 0):
                    if(pix[-1] != 0):
                        pix[-1] -= 1
                    else:
                        pix[-1] += 1

            else:
                if (pix[-1] % 2 != 0):
                    pix[-1] -= 1

            pix = tuple(pix)
            yield pix[0:3]
            yield pix[3:6]
            yield pix[6:9]

    def encode_enc(self, newimg, data):
        w = newimg.size[0]
        (x, y) = (0, 0)

        for pixel in self.modPix(newimg.getdata(), data):
            # Putting modified pixels in the new image
            newimg.putpixel((x, y), pixel)
            if (x == w - 1):
                x = 0
                y += 1
            else:
                x += 1

    # Encode data into image
    def Encode(self):
        image = Image.open(self.image_global, 'r')
        data = self.encrypt_global
        if (len(data) == 0):
            raise ValueError('Data is empty')

        newimg = image.copy()
        self.encode_enc(newimg, data)
        
        newimg.save(self.output)

    def cari(self):
        self.output = QFileDialog.getSaveFileName(None, "Save file", "Done/", "Image Files (*.png)")[0]
        self.outputname.setText(self.output)
    
    def vincky(self):
        start = timer()
        self.Encrypt()
        self.Encode()
        end = timer()
        elapsed_time = end - start
        print(f"[INFO] Waktu diperlukan pada proses penyisipan adalah {elapsed_time} s")
        msg = QMessageBox.question(self, "Penyisipan Pesan Rahasia Berhasil", "Apakah kamu ingin menyisipkan pesan rahasia lagi?", QMessageBox.No|QMessageBox.Yes)
        if (msg == QMessageBox.No):
            QApplication.closeAllWindows()
        elif(msg == QMessageBox.Yes):
            screen=Encrypt()
            widget.addWidget(screen)
            widget.setCurrentIndex(widget.currentIndex()+1)

app=QApplication(sys.argv)
mainwindow=MainWindow()
widget=QtWidgets.QStackedWidget()
widget.addWidget(mainwindow)
screen2=Decrypt()
screen3=Encrypt()
widget.setFixedWidth(800)
widget.setFixedHeight(390)
widget.show()
sys.exit(app.exec_())